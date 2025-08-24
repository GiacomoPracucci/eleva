"""
Document parsing service for extracting text from various file formats.

This module provides a unified interface for extracting text content from
different file types (PDF, DOCX, TXT, MD) while preserving structure
and metadata where possible.
"""

import io
import logging
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
import mimetypes

# Document parsing libraries
import PyPDF2
from docx import Document as DocxDocument
import markdown
import chardet

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """
    Custom exception for document parsing failures.
    
    This exception provides detailed error information that can be
    stored in the database and shown to users.
    """
    pass


class DocumentParser:
    """
    Unified document parser supporting multiple file formats.
    
    This class provides methods to extract text and metadata from various
    document formats. It's designed to be extensible for adding new formats
    and includes robust error handling for corrupted or problematic files.
    
    Supported formats:
    - PDF (.pdf)
    - Microsoft Word (.docx)
    - Plain text (.txt)
    - Markdown (.md)
    """
    
    # Maximum file size in bytes (50 MB default)
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    # Supported MIME types and their corresponding parsers
    SUPPORTED_TYPES = {
        'application/pdf': 'parse_pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'parse_docx',
        'text/plain': 'parse_text',
        'text/markdown': 'parse_markdown',
        'text/x-markdown': 'parse_markdown',
    }
    
    # File extension fallbacks for when MIME type detection fails
    EXTENSION_MAPPING = {
        '.pdf': 'parse_pdf',
        '.docx': 'parse_docx',
        '.txt': 'parse_text',
        '.md': 'parse_markdown',
        '.markdown': 'parse_markdown',
    }
    
    def __init__(self):
        """
        Initialize the document parser.
        
        Sets up any necessary configuration and validates that required
        libraries are available.
        """
        self.stats = {
            'total_parsed': 0,
            'total_failed': 0,
            'bytes_processed': 0
        }
    
    async def parse_document(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a document and extract its text content and metadata.
        
        This is the main entry point for document parsing. It automatically
        detects the file type and routes to the appropriate parser.
        
        Args:
            file_content: The raw bytes of the document
            filename: Original filename (used for type detection)
            mime_type: Optional MIME type override
            
        Returns:
            A tuple of (extracted_text, metadata_dict)
            
        Raises:
            DocumentParsingError: If parsing fails or file type is unsupported
        """
        # Validate file size
        if len(file_content) > self.MAX_FILE_SIZE:
            raise DocumentParsingError(
                f"File size ({len(file_content)} bytes) exceeds maximum "
                f"allowed size ({self.MAX_FILE_SIZE} bytes)"
            )
        
        # Detect file type
        parser_method = self._detect_parser(filename, mime_type, file_content)
        if not parser_method:
            raise DocumentParsingError(
                f"Unsupported file type for '{filename}'. "
                f"Supported types: PDF, DOCX, TXT, MD"
            )
        
        # Calculate file hash for deduplication
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Parse the document
        try:
            logger.info(f"Parsing document '{filename}' using {parser_method.__name__}")
            text, metadata = await parser_method(file_content, filename)
            
            # Add common metadata
            metadata.update({
                'original_filename': filename,
                'file_size': len(file_content),
                'file_hash': file_hash,
                'parser_used': parser_method.__name__,
            })
            
            # Update statistics
            self.stats['total_parsed'] += 1
            self.stats['bytes_processed'] += len(file_content)
            
            # Validate extracted text
            if not text or len(text.strip()) < 10:
                raise DocumentParsingError(
                    f"Document '{filename}' appears to be empty or contains insufficient text"
                )
            
            return text, metadata
            
        except DocumentParsingError:
            self.stats['total_failed'] += 1
            raise
        except Exception as e:
            self.stats['total_failed'] += 1
            logger.error(f"Unexpected error parsing '{filename}': {str(e)}")
            raise DocumentParsingError(
                f"Failed to parse document '{filename}': {str(e)}"
            )
    
    def _detect_parser(
        self,
        filename: str,
        mime_type: Optional[str],
        file_content: bytes
    ):
        """
        Detect the appropriate parser method for a file.
        
        Uses a combination of MIME type, file extension, and content
        inspection to determine the correct parser.
        
        Returns:
            The parser method to use, or None if unsupported
        """
        # Try MIME type first
        if mime_type and mime_type in self.SUPPORTED_TYPES:
            return getattr(self, self.SUPPORTED_TYPES[mime_type])
        
        # Try file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext in self.EXTENSION_MAPPING:
            return getattr(self, self.EXTENSION_MAPPING[file_ext])
        
        # Try to detect MIME type from content
        detected_mime = mimetypes.guess_type(filename)[0]
        if detected_mime and detected_mime in self.SUPPORTED_TYPES:
            return getattr(self, self.SUPPORTED_TYPES[detected_mime])
        
        # Special case: check if it's plain text by trying to decode
        try:
            file_content.decode('utf-8')
            return self.parse_text
        except UnicodeDecodeError:
            pass
        
        return None
    
    async def parse_pdf(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from PDF files.
        
        Uses PyPDF2 to extract text while preserving as much structure
        as possible. Also extracts PDF metadata like author, title, etc.
        
        Args:
            file_content: PDF file as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text_parts = []
        metadata = {
            'file_type': 'pdf',
            'pages': 0,
            'pdf_metadata': {}
        }
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract PDF metadata
            if pdf_reader.metadata:
                pdf_meta = pdf_reader.metadata
                metadata['pdf_metadata'] = {
                    'title': pdf_meta.get('/Title', ''),
                    'author': pdf_meta.get('/Author', ''),
                    'subject': pdf_meta.get('/Subject', ''),
                    'creator': pdf_meta.get('/Creator', ''),
                    'producer': pdf_meta.get('/Producer', ''),
                    'creation_date': str(pdf_meta.get('/CreationDate', '')),
                    'modification_date': str(pdf_meta.get('/ModDate', '')),
                }
            
            # Extract text from each page
            num_pages = len(pdf_reader.pages)
            metadata['pages'] = num_pages
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page marker for better chunking later
                        text_parts.append(f"\n[Page {page_num}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    text_parts.append(f"\n[Page {page_num}]\n[Text extraction failed]")
            
            if not text_parts:
                raise DocumentParsingError(
                    f"No text could be extracted from PDF '{filename}'. "
                    "The PDF might be image-based or encrypted."
                )
            
            full_text = "\n".join(text_parts)
            return full_text, metadata
            
        except PyPDF2.errors.PdfReadError as e:
            raise DocumentParsingError(f"Invalid or corrupted PDF file: {str(e)}")
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse PDF: {str(e)}")
    
    async def parse_docx(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from Microsoft Word (.docx) files.
        
        Preserves paragraph structure and extracts document properties.
        
        Args:
            file_content: DOCX file as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text_parts = []
        metadata = {
            'file_type': 'docx',
            'paragraphs': 0,
            'tables': 0,
            'document_properties': {}
        }
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # Extract document properties
            core_props = doc.core_properties
            if core_props:
                metadata['document_properties'] = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or '',
                }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    metadata['paragraphs'] += 1
            
            # Extract text from tables
            for table in doc.tables:
                metadata['tables'] += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise DocumentParsingError(
                    f"No text could be extracted from DOCX '{filename}'"
                )
            
            full_text = "\n\n".join(text_parts)
            return full_text, metadata
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}")
    
    async def parse_text(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from plain text files with encoding detection.
        
        Automatically detects character encoding to handle various text files.
        
        Args:
            file_content: Text file as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {
            'file_type': 'text',
            'encoding': 'unknown',
            'lines': 0
        }
        
        try:
            # Detect encoding
            detection = chardet.detect(file_content)
            encoding = detection.get('encoding', 'utf-8')
            confidence = detection.get('confidence', 0)
            
            metadata['encoding'] = encoding
            metadata['encoding_confidence'] = confidence
            
            # Decode text
            try:
                text = file_content.decode(encoding)
            except (UnicodeDecodeError, TypeError):
                # Fallback to UTF-8 with error handling
                text = file_content.decode('utf-8', errors='replace')
                metadata['encoding'] = 'utf-8 (fallback)'
            
            # Count lines
            metadata['lines'] = len(text.splitlines())
            
            if not text.strip():
                raise DocumentParsingError(
                    f"Text file '{filename}' appears to be empty"
                )
            
            return text, metadata
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse text file: {str(e)}")
    
    async def parse_markdown(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from Markdown files, converting to plain text.
        
        Removes Markdown syntax while preserving document structure.
        
        Args:
            file_content: Markdown file as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {
            'file_type': 'markdown',
            'has_code_blocks': False,
            'has_tables': False,
            'headers': []
        }
        
        try:
            # Decode markdown content
            md_text = file_content.decode('utf-8')
            
            # Extract some metadata before conversion
            lines = md_text.splitlines()
            for line in lines:
                if line.startswith('#'):
                    # Extract headers
                    header_level = len(line) - len(line.lstrip('#'))
                    header_text = line.lstrip('#').strip()
                    if header_text:
                        metadata['headers'].append({
                            'level': header_level,
                            'text': header_text
                        })
                elif line.strip().startswith('```'):
                    metadata['has_code_blocks'] = True
                elif '|' in line and line.count('|') >= 2:
                    metadata['has_tables'] = True
            
            # Convert to HTML then extract text (preserves structure better)
            html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
            
            # Simple HTML to text conversion (you might want to use BeautifulSoup for better results)
            import re
            text = re.sub('<[^<]+?>', '', html)
            text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
            
            if not text.strip():
                raise DocumentParsingError(
                    f"Markdown file '{filename}' appears to be empty"
                )
            
            return text, metadata
            
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse Markdown file: {str(e)}")


# Create a global parser instance
document_parser = DocumentParser()